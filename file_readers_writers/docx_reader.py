from docx import Document
import os
  
def extract_docx_text(document):
  """
  Extracts text from a docx file including formatting.

  Args:
    file_path: Path to the docx file.

  Returns:
    A string containing the extracted text with formatting.
  """
  try:
    # Handle .docx files
    if document is not None:
        # Check if the uploaded file is a .docx file
        if document.name.endswith(".docx"):
            # Read and process the .docx file using the Document function from the python-docx package
            doc = Document(document)  # Read the uploaded document
            # doc_text = "\n".join([para.text for para in doc.paragraphs])  # Extract all text from the document
            # st.write("Document content (DOCX):")
            # st.write(doc_text)
            fullText = []
            for para in doc.paragraphs:
              fullText.append(para.text)
            return '\n'.join(fullText)
  except Exception as e:
    print(f"Error extracting text from docx file: {e}")
    return None

def extract_folder_and_filename(filepath):
  """
  Extracts the folder path and filename from a provided filepath string.

  Args:
    filepath: The path to the file.

  Returns:
    A tuple containing the folder path and filename.
  """
  import os
  folder_path = os.path.dirname(filepath)
  filename = os.path.basename(filepath)
  return folder_path, filename

def get_folder_path(filepath):
  """
  Retrieves the folder path from an absolute filepath.

  Args:
    filepath: The absolute path to the file.

  Returns:
    The folder path.
  """
  return os.path.dirname(filepath)